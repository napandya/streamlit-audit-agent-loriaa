"""
Tests for ingestion.parsers sub-package.
"""
import io
import tempfile
import os
import pytest
import pandas as pd

from ingestion.parsers.csv_parser import parse_csv
from ingestion.parsers.pdf_parser import parse_pdf
from ingestion.parsers.docx_parser import parse_docx
from ingestion.parsers import detect_document_type, ParsedDocument


# ---------------------------------------------------------------------------
# CSV parser tests
# ---------------------------------------------------------------------------

def test_csv_parser_rent_roll(sample_rent_roll_csv_path):
    """CSV parser correctly parses rent roll and detects document type."""
    doc = parse_csv(sample_rent_roll_csv_path)
    assert doc.dataframe is not None
    assert not doc.dataframe.empty
    assert doc.file_type == "csv"
    assert doc.document_type == "rent_roll"
    # Should contain at least unit/status columns
    cols_lower = [c.lower() for c in doc.dataframe.columns]
    assert any("unit" in c or "status" in c for c in cols_lower)


def test_csv_parser_projection(sample_projection_csv_path):
    """CSV parser detects projection document type."""
    doc = parse_csv(sample_projection_csv_path)
    assert doc.dataframe is not None
    assert doc.document_type == "projection"


def test_csv_parser_empty_file():
    """CSV parser with an empty file returns a doc without crashing."""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
        f.write("")
        tmp_path = f.name
    try:
        doc = parse_csv(tmp_path)
        assert isinstance(doc, ParsedDocument)
    finally:
        os.unlink(tmp_path)


def test_csv_parser_minimal():
    """CSV parser with a minimal valid CSV works correctly."""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False, encoding="utf-8") as f:
        f.write("unit,status,rent\n101,C,1250\n102,NTV,1250\n")
        tmp_path = f.name
    try:
        doc = parse_csv(tmp_path)
        assert doc.dataframe is not None
        assert len(doc.dataframe) == 2
    finally:
        os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# PDF parser tests
# ---------------------------------------------------------------------------

def test_pdf_parser_basic(sample_pdf_path):
    """PDF parser extracts non-empty text and sets file_type."""
    doc = parse_pdf(sample_pdf_path)
    assert doc.file_type == "pdf"
    assert doc.raw_text.strip() != ""


def test_pdf_parser_document_type(sample_pdf_path):
    """PDF parser correctly detects projection document type from filename/content."""
    doc = parse_pdf(sample_pdf_path)
    # The PDF filename contains 'projection'
    assert doc.document_type == "projection"


# ---------------------------------------------------------------------------
# DOCX parser tests
# ---------------------------------------------------------------------------

def test_docx_parser_programmatic():
    """DOCX parser extracts text from a programmatically created docx."""
    from docx import Document

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "test_doc.docx")
        doc = Document()
        doc.add_paragraph("Unit 101 - Rent Roll Summary")
        doc.add_paragraph("Status: C, Rent: $1,250")
        doc.save(path)

        result = parse_docx(path)
        assert result.file_type == "docx"
        assert "Unit 101" in result.raw_text
        assert "Rent Roll" in result.raw_text


# ---------------------------------------------------------------------------
# Document type detection tests
# ---------------------------------------------------------------------------

def test_detect_document_type_rent_roll():
    assert detect_document_type("Rent Roll Q1.csv") == "rent_roll"


def test_detect_document_type_projection():
    assert detect_document_type("Recurring Transaction Projection.pdf") == "projection"


def test_detect_document_type_concession():
    assert detect_document_type("concession_schedule.xlsx") == "concession"


def test_detect_document_type_unknown():
    assert detect_document_type("random_file.csv", "no relevant content") == "unknown"
