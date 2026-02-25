"""
DOCX parser — returns a ParsedDocument.
"""
from pathlib import Path

from ingestion.parsers import ParsedDocument, detect_document_type


def parse_docx(file_path: str) -> ParsedDocument:
    """
    Parse a Word (.docx) file using python-docx and return a ParsedDocument.
    Extracts all paragraphs and table cell text as raw_text.
    """
    from docx import Document  # python-docx

    path = Path(file_path)
    doc = Document(str(path))

    text_parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    raw_text = "\n".join(text_parts)
    doc_type = detect_document_type(path.name, raw_text[:2000])

    return ParsedDocument(
        file_name=path.name,
        file_type="docx",
        raw_text=raw_text,
        dataframe=None,
        document_type=doc_type,
    )
